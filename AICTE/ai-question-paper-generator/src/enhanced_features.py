"""
Enhanced Question Paper Generator Module

This module extends the existing system to support:
- Unit-based question selection
- Total marks-based distribution
- Word document (.docx) export
- PDF and Word input parsing
"""

import random
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class EnhancedQuestionSelector:
    """Enhanced question selector with unit-based and marks-based selection"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.questions = []
        
    def load_questions(self, questions: List[Dict[str, Any]]):
        """Load questions into the selector"""
        self.questions = questions
        
    def get_available_units(self) -> List[str]:
        """Get list of available units/topics from loaded questions"""
        units = set()
        for question in self.questions:
            # Check multiple possible fields for unit/topic information
            unit = question.get('unit') or question.get('topic') or question.get('subject')
            if unit:
                units.add(str(unit))
        return sorted(list(units))
    
    def select_questions_by_units_and_marks(
        self, 
        selected_units: List[str], 
        total_marks: int,
        marks_distribution: Optional[Dict[str, int]] = None
    ) -> Dict[str, Any]:
        """
        Select questions based on units and total marks
        
        Args:
            selected_units: List of units to include
            total_marks: Total marks for the question paper
            marks_distribution: Optional custom distribution (e.g., {'2': 20, '16': 64})
            
        Returns:
            Dictionary with selected questions and configuration
        """
        
        # Filter questions by selected units
        unit_questions = []
        for question in self.questions:
            question_unit = question.get('unit') or question.get('topic') or question.get('subject')
            if question_unit and str(question_unit) in selected_units:
                unit_questions.append(question)
        
        if not unit_questions:
            raise ValueError("No questions found for the selected units")
        
        # Calculate optimal marks distribution if not provided
        if not marks_distribution:
            marks_distribution = self._calculate_optimal_distribution(total_marks)
        
        # Separate questions by marks
        questions_by_marks = {}
        for question in unit_questions:
            marks = int(question.get('marks', 2))  # Default to 2 marks
            if marks not in questions_by_marks:
                questions_by_marks[marks] = []
            questions_by_marks[marks].append(question)
        
        # Select questions according to distribution
        selected_questions = []
        actual_marks = 0
        selection_summary = {}
        
        for marks_value, required_count in marks_distribution.items():
            marks_key = int(marks_value)
            if marks_key in questions_by_marks:
                available_questions = questions_by_marks[marks_key]
                
                # Randomly select required number of questions
                if len(available_questions) >= required_count:
                    selected = random.sample(available_questions, required_count)
                else:
                    selected = available_questions  # Take all available
                    self.logger.warning(f"Only {len(available_questions)} questions available for {marks_key} marks, needed {required_count}")
                
                selected_questions.extend(selected)
                actual_marks += marks_key * len(selected)
                selection_summary[f"{marks_key}_marks"] = len(selected)
        
        # Calculate choice options for 16-mark questions
        sixteen_mark_count = selection_summary.get('16_marks', 0)
        choice_options = 2 if sixteen_mark_count > 0 else 0
        
        return {
            'questions': selected_questions,
            'total_marks': actual_marks,
            'distribution': selection_summary,
            'choice_options': choice_options,
            'units_covered': selected_units
        }
    
    def _calculate_optimal_distribution(self, total_marks: int) -> Dict[str, int]:
        """Calculate optimal distribution of 2-mark and 16-mark questions"""
        
        # Common exam patterns
        if total_marks <= 50:
            # Small test: mostly 2-mark questions
            two_mark_count = min(total_marks // 2, 20)
            remaining_marks = total_marks - (two_mark_count * 2)
            sixteen_mark_count = max(0, remaining_marks // 16)
        elif total_marks <= 100:
            # Standard exam: balanced approach
            sixteen_mark_count = min(total_marks // 20, 4)  # ~20% in 16-mark questions
            remaining_marks = total_marks - (sixteen_mark_count * 16)
            two_mark_count = remaining_marks // 2
        else:
            # Large exam: more 16-mark questions
            sixteen_mark_count = min(total_marks // 16, 6)
            remaining_marks = total_marks - (sixteen_mark_count * 16)
            two_mark_count = remaining_marks // 2
        
        return {
            '2': max(1, two_mark_count),
            '16': max(0, sixteen_mark_count)
        }


class WordDocumentGenerator:
    """Generator for Word document question papers"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Please install with: pip install python-docx")
    
    def generate_question_paper(
        self,
        questions: List[Dict[str, Any]],
        output_path: str,
        paper_config: Dict[str, Any]
    ) -> bool:
        """
        Generate a Word document question paper
        
        Args:
            questions: List of question dictionaries
            output_path: Path for output .docx file
            paper_config: Configuration dictionary
            
        Returns:
            bool: Success status
        """
        try:
            doc = Document()
            
            # Set up document properties
            section = doc.sections[0]
            section.page_height = Inches(11.7)  # A4 height
            section.page_width = Inches(8.3)    # A4 width
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            
            # Add header
            self._add_header(doc, paper_config)
            
            # Separate questions by marks
            two_mark_questions = [q for q in questions if int(q.get('marks', 2)) == 2]
            sixteen_mark_questions = [q for q in questions if int(q.get('marks', 2)) == 16]
            
            # Add sections
            if two_mark_questions:
                self._add_section_a(doc, two_mark_questions, paper_config)
            
            if sixteen_mark_questions:
                self._add_section_b(doc, sixteen_mark_questions, paper_config)
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"Word document saved to: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating Word document: {e}")
            return False
    
    def _add_header(self, doc, config: Dict[str, Any]):
        """Add header section to the document"""
        
        # Title
        title = doc.add_heading(config.get('title', 'Question Paper'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Exam details
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subject = config.get('subject', 'Subject Name')
        duration = config.get('duration', '3 Hours')
        max_marks = config.get('total_marks', 100)
        
        details_para.add_run(f"Subject: {subject}").bold = True
        details_para.add_run(f"\t\t\tDuration: {duration}\n")
        details_para.add_run(f"Maximum Marks: {max_marks}").bold = True
        details_para.add_run(f"\t\t\tDate: _____________")
        
        # Instructions
        doc.add_paragraph()
        instructions = doc.add_paragraph("INSTRUCTIONS:")
        instructions.runs[0].bold = True
        
        doc.add_paragraph("1. Read all questions carefully before answering.")
        doc.add_paragraph("2. Write your answers in the space provided.")
        doc.add_paragraph("3. All questions are compulsory unless otherwise specified.")
        
        doc.add_paragraph("\n" + "="*80)
    
    def _add_section_a(self, doc, questions: List[Dict], config: Dict):
        """Add Section A (2-mark questions)"""
        
        section_title = doc.add_heading("SECTION A - Short Answer Questions (2 Marks Each)", level=1)
        
        count = len(questions)
        instruction = f"Answer any {count} questions. Each question carries 2 marks."
        doc.add_paragraph(instruction).runs[0].italic = True
        
        doc.add_paragraph()
        
        for i, question in enumerate(questions, 1):
            # Question
            q_para = doc.add_paragraph()
            q_para.add_run(f"Q{i}. ").bold = True
            q_para.add_run(f"{question.get('question', 'N/A')} ")
            q_para.add_run("[2 marks]").bold = True
            
            # Answer space
            doc.add_paragraph("_" * 80)
            doc.add_paragraph()
    
    def _add_section_b(self, doc, questions: List[Dict], config: Dict):
        """Add Section B (16-mark questions with choices)"""
        
        doc.add_paragraph("\n" + "="*80)
        section_title = doc.add_heading("SECTION B - Long Answer Questions (16 Marks Each)", level=1)
        
        choice_options = config.get('choice_options', 2)
        question_count = len(questions) // choice_options
        
        instruction = f"Answer any {question_count} questions. Each question carries 16 marks. Choose from the given options."
        doc.add_paragraph(instruction).runs[0].italic = True
        
        doc.add_paragraph()
        
        # Group questions into choices
        for i in range(0, len(questions), choice_options):
            question_num = (i // choice_options) + 1
            
            # Question header
            q_header = doc.add_paragraph()
            q_header.add_run(f"Question {question_num}: (Choose any one)").bold = True
            
            # Choice options
            for j in range(choice_options):
                if i + j < len(questions):
                    question = questions[i + j]
                    option_label = chr(ord('a') + j)  # a, b, c, etc.
                    
                    choice_para = doc.add_paragraph()
                    choice_para.add_run(f"{option_label}) ").bold = True
                    choice_para.add_run(f"{question.get('question', 'N/A')} ")
                    choice_para.add_run("[16 marks]").bold = True
            
            # Answer space
            doc.add_paragraph("\nAnswer space for chosen option:")
            for _ in range(15):  # Multiple lines for long answers
                doc.add_paragraph("_" * 80)
            
            doc.add_paragraph()


class EnhancedInputParser:
    """Enhanced parser for PDF and Word input files"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse_pdf_questions(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse questions from PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF libraries not available. Please install PyPDF2 and pdfplumber")
        
        questions = []
        try:
            with open(file_path, 'rb') as file:
                # Try pdfplumber first (better text extraction)
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text = ""
                        for page in pdf.pages:
                            text += page.extract_text() + "\n"
                except:
                    # Fallback to PyPDF2
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            
            # Basic question parsing (this can be enhanced based on PDF format)
            questions = self._parse_text_to_questions(text)
            
        except Exception as e:
            self.logger.error(f"Error parsing PDF: {e}")
            
        return questions
    
    def parse_docx_questions(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse questions from Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available")
        
        questions = []
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            questions = self._parse_text_to_questions(text)
            
        except Exception as e:
            self.logger.error(f"Error parsing Word document: {e}")
            
        return questions
    
    def _parse_text_to_questions(self, text: str) -> List[Dict[str, Any]]:
        """Parse text content to extract questions"""
        questions = []
        lines = text.split('\n')
        
        current_question = None
        question_id = 1
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for question patterns (this is a basic implementation)
            if any(line.startswith(prefix) for prefix in ['Q', 'Question', 'q', 'question', str(question_id)]):
                # Save previous question
                if current_question:
                    questions.append(current_question)
                
                # Start new question
                current_question = {
                    'id': question_id,
                    'question': line,
                    'topic': 'General',
                    'difficulty': 'medium',
                    'type': 'text',
                    'marks': 2  # Default marks
                }
                question_id += 1
            
            elif current_question and line:
                # Continue previous question
                current_question['question'] += " " + line
        
        # Don't forget the last question
        if current_question:
            questions.append(current_question)
        
        return questions
